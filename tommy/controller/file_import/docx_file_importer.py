import os
from docx import Document
from tommy.controller.file_import import file_importer_base
from tommy.controller.file_import.metadata import Metadata
from tommy.controller.file_import.raw_body import RawBody
from tommy.controller.file_import.raw_file import RawFile


class WordFileImporter(file_importer_base.FileImporterBase):
    """
    Handles importing of Word files
    """

    def __init__(self) -> None:
        """
        Initializes a new instance of the class.
        """
        pass

    def compatible_file(self, path: str) -> bool:
        """
        A Word file is compatible with this parser if and only if
        it can be opened and read without errors.

        :param path: The string path to the Word file to be checked for
                     compatibility.
        :return: bool: True if the file is compatible, False otherwise.
        """
        try:
            Document(path)
            return True
        except Exception as e:
            print(f"Error reading file '{path}': {e}")
            return False

    def load_file(self, path: str) -> RawFile:
        """
        Loads a Word file and returns a File object.

        :param path: The string path to the Word file.
        :return: File: A File object generated from the Word file.
        """
        doc = Document(path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

        return self.generate_file(text, path)

    def generate_file(self, text: str, path) -> RawFile:
        """
        Generates a File object from a Word file.

        :param text: A string representing the text of the Word file.
        :param path: The string path to the Word file.
        :return: A RawFile object generated from the Word file
        containing metadata and the raw text of the file.
        """
        return RawFile(
            metadata=Metadata(author=None,
                              title=None, date=None,
                              url=None, path=path,
                              format="docx",
                              length=len(text.split(" ")),
                              name=os.path.relpath(path).split(".")[0],
                              size=stat(path).st_size),
            body=RawBody(body=text))